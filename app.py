import streamlit as st
import json
import re
import os
import io
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="SimGroupAI",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded",
)

PRIORITIES = [
    "Price & Value", "Product Features & Functionality", "Design & Aesthetics",
    "Ease of Use & UX", "Brand Trust & Credibility", "Privacy & Data Security",
    "Performance & Quality", "Sustainability & Ethics", "Health & Safety",
    "Innovation & Novelty", "Social Status & Prestige", "Convenience & Speed",
    "Customer Support", "Accessibility & Inclusivity", "Environmental Impact",
    "Community & Social Impact", "Family-Friendliness", "Durability & Reliability",
]

# ─────────────────────────────────────────────
# CSS
# ─────────────────────────────────────────────
def inject_css():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700&display=swap');

    html, body, [class*="css"] { font-family: 'Roboto', sans-serif; font-size: 14px; }

    .stApp { background: #1c1c1e; color: #e8eaed; }

    /* ── Sidebar ── */
    section[data-testid="stSidebar"] {
        background: #28282b !important;
        border-right: 1px solid #3c3c3f;
        display: block !important;
        visibility: visible !important;
    }
    section[data-testid="stSidebar"] label,
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] span,
    section[data-testid="stSidebar"] div { color: #bdc1c6 !important; }
    section[data-testid="stSidebar"] .stMarkdown h3 { color: #e8eaed !important; }
    [data-testid="stSidebarHeader"] { display: none !important; }
    [data-testid="collapsedControl"] {
        display: flex !important; visibility: visible !important; opacity: 1 !important;
        pointer-events: auto !important; position: fixed !important;
        top: 70px !important; left: 0 !important; z-index: 999999 !important;
        background: #1a73e8 !important; border: none !important;
        border-radius: 0 6px 6px 0 !important; padding: 10px 8px !important;
        color: #fff !important; cursor: pointer !important;
    }

    /* ── Top nav bar ── */
    .ga-topbar {
        background: #28282b; border-bottom: 1px solid #3c3c3f;
        padding: 0 24px; margin: -1rem -1rem 28px -1rem;
        display: flex; align-items: center; height: 56px; gap: 16px;
    }
    .ga-logo {
        font-size: 1.15rem; font-weight: 700; color: #e8eaed;
        display: flex; align-items: center; gap: 8px;
    }
    .ga-logo-dot { width: 8px; height: 8px; border-radius: 50%; background: #1a73e8; display: inline-block; }
    .ga-property { font-size: 0.78rem; color: #9aa0a6; border-left: 1px solid #3c3c3f; padding-left: 14px; }

    /* ── Scorecards ── */
    .ga-scorecard-row { display: flex; gap: 16px; margin-bottom: 24px; flex-wrap: wrap; }
    .ga-scorecard {
        background: #28282b; border: 1px solid #3c3c3f; border-radius: 8px;
        padding: 16px 20px; flex: 1; min-width: 140px;
    }
    .ga-scorecard-label { font-size: 0.72rem; color: #9aa0a6; text-transform: uppercase; letter-spacing: 0.7px; margin-bottom: 6px; }
    .ga-scorecard-value { font-size: 2rem; font-weight: 400; color: #e8eaed; line-height: 1; }
    .ga-scorecard-delta { font-size: 0.75rem; margin-top: 4px; }
    .delta-green { color: #34a853; } .delta-yellow { color: #fbbc04; } .delta-red { color: #ea4335; }

    /* ── Progress bars — ALL variants need height:100% ── */
    .ga-bar-row { margin-bottom: 10px; }
    .ga-bar-label-row { display: flex; justify-content: space-between; font-size: 0.78rem; color: #bdc1c6; margin-bottom: 4px; }
    .ga-bar-track { background: #3c3c3f; border-radius: 3px; height: 5px; overflow: hidden; }
    .ga-bar-fill, .ga-bar-fill-warn, .ga-bar-fill-danger, .ga-bar-fill-green {
        height: 100%; border-radius: 3px;
    }
    .ga-bar-fill       { background: #1a73e8; }
    .ga-bar-fill-warn  { background: #fbbc04; }
    .ga-bar-fill-danger{ background: #ea4335; }
    .ga-bar-fill-green { background: #34a853; }

    /* ── Findings rows ── */
    .ga-finding-row {
        display: flex; align-items: flex-start; gap: 10px;
        padding: 8px 0; border-bottom: 1px solid #3c3c3f;
        font-size: 0.84rem; color: #bdc1c6; line-height: 1.45;
    }
    .ga-finding-row:last-child { border-bottom: none; }
    .ga-finding-dot { width: 6px; height: 6px; border-radius: 50%; margin-top: 6px; flex-shrink: 0; }
    .dot-red { background: #ea4335; } .dot-green { background: #34a853; } .dot-yellow { background: #fbbc04; }

    /* ── Verdict badge ── */
    .verdict-badge { display: inline-block; border-radius: 4px; padding: 4px 12px; font-size: 0.78rem; font-weight: 500; }
    .verdict-green  { background: #1e3a2a; color: #81c995; border: 1px solid #34a853; }
    .verdict-yellow { background: #3a3010; color: #fdd663; border: 1px solid #fbbc04; }
    .verdict-orange { background: #3a2810; color: #ffb74d; border: 1px solid #f57c00; }
    .verdict-red    { background: #3a1e1e; color: #f28b82; border: 1px solid #ea4335; }

    /* ── Honesty chip ── */
    .honesty-chip { font-size: 0.68rem; padding: 2px 7px; border-radius: 10px; font-weight: 500; margin-left: 6px; vertical-align: middle; }
    .hc-high { background: #1e3a2a; color: #81c995; }
    .hc-mid  { background: #3a3010; color: #fdd663; }
    .hc-low  { background: #3a1e1e; color: #f28b82; }

    /* ── Transcript bubbles ── */
    .tx-bubble {
        padding: 10px 14px; border-radius: 6px; margin: 5px 0;
        background: #28282b; border: 1px solid #3c3c3f;
        font-size: 0.87rem; line-height: 1.55; color: #bdc1c6;
    }
    .tx-bubble:hover { border-color: #5f6368; }
    .tx-name { font-size: 0.72rem; font-weight: 500; color: #8ab4f8; margin-bottom: 4px; text-transform: uppercase; letter-spacing: 0.5px; }
    .tx-bubble-user { background: #1a2535; border-color: #1a73e8; }
    .tx-name-user { color: #fbbc04; }

    /* ── Next step cards ── */
    .ns-card { background: #28282b; border: 1px solid #3c3c3f; border-left: 3px solid #1a73e8; border-radius: 6px; padding: 14px 18px; margin-bottom: 10px; }
    .ns-title { font-size: 0.84rem; font-weight: 500; color: #8ab4f8; margin-bottom: 4px; }
    .ns-body  { font-size: 0.81rem; color: #9aa0a6; line-height: 1.5; }
    .ns-profile { font-size: 0.73rem; color: #5f6368; margin-top: 5px; }

    /* ── Section headers ── */
    .ga-section-title {
        font-size: 0.72rem; font-weight: 500; color: #9aa0a6;
        text-transform: uppercase; letter-spacing: 0.8px;
        margin: 24px 0 12px 0; padding-bottom: 6px; border-bottom: 1px solid #3c3c3f;
        display: flex; align-items: center; gap: 6px;
    }

    /* ── CSS Tooltip ── */
    .tip-wrap { position: relative; display: inline-flex; align-items: center; cursor: default; }
    .tip-icon {
        width: 15px; height: 15px; border-radius: 50%;
        background: #3c3c3f; color: #9aa0a6;
        font-size: 0.65rem; text-align: center; line-height: 15px;
        font-weight: 700; display: inline-block; margin-left: 5px;
        border: 1px solid #5f6368;
    }
    .tip-text {
        visibility: hidden; opacity: 0;
        background: #35363a; color: #e8eaed; border: 1px solid #5f6368;
        text-align: left; border-radius: 6px; padding: 8px 12px;
        position: absolute; z-index: 99999;
        bottom: 130%; left: 50%; transform: translateX(-50%);
        width: 220px; font-size: 0.75rem; line-height: 1.45;
        font-weight: 400; text-transform: none; letter-spacing: 0;
        box-shadow: 0 4px 16px rgba(0,0,0,0.4);
        transition: opacity 0.15s ease;
        pointer-events: none;
    }
    .tip-text::after {
        content: ""; position: absolute; top: 100%; left: 50%;
        margin-left: -5px; border-width: 5px;
        border-style: solid; border-color: #5f6368 transparent transparent transparent;
    }
    .tip-wrap:hover .tip-text { visibility: visible; opacity: 1; }

    /* ── Phase nav buttons ── */
    .phase-nav { display: flex; gap: 4px; margin-bottom: 20px; }
    .phase-nav-btn {
        padding: 8px 18px; border-radius: 6px 6px 0 0;
        font-size: 0.84rem; font-weight: 500; cursor: pointer;
        border: 1px solid #3c3c3f; border-bottom: none;
        background: #28282b; color: #9aa0a6;
    }
    .phase-nav-btn.active { background: #35363a; color: #e8eaed; border-color: #5f6368; }

    /* ── Phase 2 info box ── */
    .p2-info {
        background: #1a2535; border: 1px solid #1a73e8; border-radius: 6px;
        padding: 12px 16px; font-size: 0.82rem; color: #8ab4f8;
        margin-bottom: 20px; line-height: 1.5;
    }

    /* ── Inputs ── */
    .stTextArea textarea, .stTextInput input {
        background: #28282b !important; border: 1px solid #3c3c3f !important;
        color: #e8eaed !important; border-radius: 6px !important; font-size: 0.87rem !important;
    }
    .stTextArea textarea:focus, .stTextInput input:focus {
        border-color: #1a73e8 !important; box-shadow: 0 0 0 2px rgba(26,115,232,0.2) !important;
    }

    /* ── Buttons ── */
    .stButton > button {
        background: #1a73e8 !important; color: #fff !important; font-weight: 500 !important;
        border: none !important; border-radius: 4px !important;
        padding: 8px 22px !important; font-size: 0.87rem !important;
    }
    .stButton > button:hover { background: #1557b0 !important; }

    /* ── Expander ── */
    .streamlit-expanderHeader {
        background: #28282b !important; border: 1px solid #3c3c3f !important;
        border-radius: 6px !important; color: #bdc1c6 !important; font-size: 0.84rem !important;
    }

    /* ── Tabs ── */
    .stTabs [data-baseweb="tab-list"] { background: #28282b; border-radius: 6px; padding: 3px; gap: 2px; border: 1px solid #3c3c3f; }
    .stTabs [data-baseweb="tab"] { background: transparent; color: #9aa0a6 !important; border-radius: 4px; font-size: 0.84rem; font-weight: 500; }
    .stTabs [aria-selected="true"] { background: #35363a !important; color: #e8eaed !important; }

    /* ── Misc ── */
    #MainMenu { visibility: hidden; } footer { visibility: hidden; }
    [data-testid="stDeployButton"] { display: none !important; }
    [data-testid="stStatusWidget"] { display: none !important; }
    hr { border-color: #3c3c3f !important; }
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-track { background: #1c1c1e; }
    ::-webkit-scrollbar-thumb { background: #3c3c3f; border-radius: 3px; }
    </style>
    """, unsafe_allow_html=True)


# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
def init_session_state():
    defaults = {
        "api_key": "", "simulation_done": False,
        "transcript_lines": [], "insights": {}, "auth": {},
        "focus_guide": "", "next_steps": [], "phase2_messages": [],
        "saved_panels": {}, "custom_panel": None, "panel_mode": "library",
        "selected_personas": [], "pitch": "", "selected_priorities": [],
        "raw_transcript": "", "insights_text": "",
        "active_view": "phase1",   # "phase1" or "phase2"
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ─────────────────────────────────────────────
# DATA
# ─────────────────────────────────────────────
@st.cache_data
def load_personas():
    path = os.path.join(os.path.dirname(__file__), "personas.json")
    with open(path, "r") as f:
        return json.load(f)


def get_client():
    key = st.session_state.get("api_key", "").strip()
    return OpenAI(api_key=key) if key else None


# ─────────────────────────────────────────────
# LLM CALLS
# ─────────────────────────────────────────────
def generate_custom_panel(client, description):
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": f"""Create exactly 4 diverse consumer personas for this target market: "{description}"

Return a JSON object with key "personas" containing an array of 4 objects, each with:
name, age_group, location_type, income_bracket, core_values (array), communication_style, icon (emoji), background, rating_style (Very Tough/Tough/Balanced/Generous)

Make them diverse in age, income, background, and values."""}],
        response_format={"type": "json_object"}, temperature=0.9,
    )
    parsed = json.loads(resp.choices[0].message.content.strip())
    if isinstance(parsed, list):
        return parsed
    for v in parsed.values():
        if isinstance(v, list):
            return v
    return []


def generate_simulation(client, pitch, personas, priorities):
    priority_str = ", ".join(priorities) if priorities else "general consumer value"
    persona_block = "\n".join([
        f"- {p['icon']} {p['name']} ({p['age_group']}, {p['location_type']}, {p['income_bracket']}): "
        f"{p['communication_style']} Rating style: {p.get('rating_style','Balanced')}."
        for p in personas
    ])
    system = f"""Simulate a realistic, unscripted consumer focus group with {len(personas)} participants.

PARTICIPANTS:
{persona_block}

RULES:
1. Exactly 30-35 total exchanges. Each = one participant speaking.
2. NO moderator. Participants speak freely and react to each other.
3. MANDATORY: At least 2 participants openly disagree with another.
4. MANDATORY: At least 2 express genuine skepticism or concern.
5. MANDATORY: Each participant speaks at least 4 times.
6. Very Tough/Tough personas are harder to please. Generous personas more open.
7. 2-4 sentences per response. No speeches. NO groupthink.
8. Naturally explore: {priority_str}

FORMAT (every line):
[NAME]: [message]

Start immediately. No intro text."""
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system", "content": system},
                  {"role": "user", "content": f"PITCH:\n{pitch}\n\nBegin."}],
        temperature=0.85, max_tokens=3500,
    )
    return resp.choices[0].message.content.strip()


def generate_insights(client, pitch, transcript, personas, priorities):
    priority_str = ", ".join(priorities) if priorities else "overall consumer value"
    names = [p["name"] for p in personas]
    # Use varied example scores so LLM understands these must differ per entry
    prio_example = [74, 61, 88, 55, 79]
    prio_list = priorities if priorities else ["Overall Value"]
    prio_lines = "\n".join([f"- {p}: {prio_example[i % len(prio_example)]}" for i, p in enumerate(prio_list)])
    honesty_example = [82, 71, 65, 88, 74, 59, 91, 77]
    honesty_lines = "\n".join([f"- {n}: {honesty_example[i % len(honesty_example)]}" for i, n in enumerate(names)])
    user = f"""PITCH: {pitch}
TRANSCRIPT:
{transcript}
PRIORITY DIMENSIONS: {priority_str}
PERSONAS: {json.dumps([{"name": p["name"], "rating_style": p.get("rating_style","Balanced")} for p in personas])}

Return the analysis in this EXACT structure (replace example numbers with real values — no brackets):
ADOPTION_SCORE: 74
SENTIMENT_BREAKDOWN: Positive 60% / Neutral 25% / Negative 15%
TOP_PAIN_POINTS:
- first pain point here
- second pain point here
- third pain point here
TOP_STRENGTHS:
- first strength here
- second strength here
- third strength here
KEY_OBJECTIONS:
- first objection here
- second objection here
PRIORITY_SCORES:
{prio_lines}
PERSONA_HONESTY:
{honesty_lines}
OVERALL_VERDICT: Your 1-2 sentence summary here.

Now return the ACTUAL analysis using real values from the transcript.
IMPORTANT: PRIORITY_SCORES and PERSONA_HONESTY scores must each be DIFFERENT numbers reflecting actual variation — never return the same score for every entry."""
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a senior consumer insights analyst. Return structured data in the EXACT format shown. No extra commentary, no brackets around numbers. Every persona honesty score and every priority dimension score must be a distinct number reflecting real differences — never repeat the same score for all entries."},
            {"role": "user", "content": user}
        ],
        temperature=0.3, max_tokens=1200,
    )
    return resp.choices[0].message.content.strip()


def check_authenticity(client, pitch, transcript):
    user = f"""PITCH: {pitch}
TRANSCRIPT:
{transcript}

Return in this EXACT format (replace example values):
AUTHENTICITY_SCORE: 82
GROUPTHINK_DETECTED: No
GROUPTHINK_NOTES: None
OUTLIER_VOICES:
- Name here
CORPORATE_SYCOPHANCY_RISK: Low
FILTER_VERDICT: Your 1-2 sentence verdict here."""
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a focus group quality auditor. Return EXACT format only, no brackets around numbers."},
            {"role": "user", "content": user}
        ],
        temperature=0.2, max_tokens=500,
    )
    return resp.choices[0].message.content.strip()


def generate_focus_guide(client, pitch, transcript, insights_text):
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a qualitative research consultant. Write a practical human focus group guide."},
            {"role": "user", "content": f"""PITCH: {pitch}
INSIGHTS: {insights_text}
TRANSCRIPT EXCERPT: {transcript[:2000]}

Write a Human Focus Group Guide with:
1. RESEARCH OBJECTIVES (3 points)
2. SCREENING CRITERIA (4-5 criteria)
3. WARM-UP QUESTIONS (2)
4. CORE DISCUSSION QUESTIONS (5-6 based on unresolved tensions)
5. PROBING FOLLOW-UPS (3)
6. CLOSING EXERCISE (1 activity)"""}
        ],
        temperature=0.5, max_tokens=1000,
    )
    return resp.choices[0].message.content.strip()


def generate_next_steps(client, pitch, insights_text, personas):
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a consumer research strategist."},
            {"role": "user", "content": f"""PITCH: {pitch}
FINDINGS: {insights_text}
CURRENT PANEL: {json.dumps([{"name": p["name"], "age_group": p["age_group"]} for p in personas])}

Recommend exactly 3 follow-up research populations:

PANEL_1_TITLE: name
PANEL_1_RATIONALE: 1-2 sentences
PANEL_1_PROFILE: age, location, income, key trait

PANEL_2_TITLE: name
PANEL_2_RATIONALE: 1-2 sentences
PANEL_2_PROFILE: age, location, income, key trait

PANEL_3_TITLE: name
PANEL_3_RATIONALE: 1-2 sentences
PANEL_3_PROFILE: age, location, income, key trait"""}
        ],
        temperature=0.6, max_tokens=700,
    )
    return resp.choices[0].message.content.strip()


def generate_followup(client, pitch, original_transcript, personas, phase2_history, user_question):
    persona_block = "\n".join([f"- {p['icon']} {p['name']} ({p['age_group']}): {p['communication_style']}" for p in personas])
    history_block = "\n".join([
        f"{'Moderator' if m['role'] == 'user' else m.get('speaker','Panel')}: {m['content']}"
        for m in phase2_history[-10:]
    ])
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": f"""You are simulating a follow-up Q&A with focus group participants.

ORIGINAL PITCH: {pitch}
PARTICIPANTS:
{persona_block}

RULES:
1. Select 2-3 most relevant participants to respond.
2. If question names a specific participant, ONLY that person responds.
3. 2-3 sentences each. Conversational. May reference the original session.
4. Maintain each persona's voice and values.

FORMAT:
[NAME]: [response]

Only participant responses. No narration."""},
            {"role": "user", "content": f"ORIGINAL TRANSCRIPT:\n{original_transcript}\n\nCONVERSATION SO FAR:\n{history_block}\n\nMODERATOR: {user_question}"},
        ],
        temperature=0.8, max_tokens=600,
    )
    return resp.choices[0].message.content.strip()


# ─────────────────────────────────────────────
# PARSERS
# ─────────────────────────────────────────────
def parse_transcript(text, personas):
    lines = []
    persona_map = {p["name"].lower(): p for p in personas}
    for line in text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        m = re.match(r"^\[?([^\]:\n]+)\]?:\s*(.+)$", line)
        if m:
            speaker = m.group(1).strip()
            message = m.group(2).strip()
            persona = None
            for key, p in persona_map.items():
                if key in speaker.lower() or speaker.lower() in key:
                    persona = p
                    break
            lines.append({"speaker": speaker, "message": message, "persona": persona})
    return lines


def parse_insights(text):
    result = {}

    def extract_num(key):
        """First integer after KEY:, handles plain 72 or bracketed [72]."""
        m = re.search(rf"{key}[^\d]*(\d+)", text, re.IGNORECASE)
        return m.group(1).strip() if m else "N/A"

    def extract_line(key):
        """Rest of the line after KEY:"""
        m = re.search(rf"{key}\s*:\s*(.+?)$", text, re.IGNORECASE | re.MULTILINE)
        return m.group(1).strip() if m else "N/A"

    def extract_bullets(start_key, stop_key=None):
        """Bullet lines between start_key and stop_key (or next all-caps section)."""
        if stop_key:
            m = re.search(
                rf"{start_key}\s*:.*?\n(.*?)(?={stop_key}\s*:|$)",
                text, re.IGNORECASE | re.DOTALL
            )
        else:
            m = re.search(rf"{start_key}\s*:.*?\n((?:\s*[-•*].+\n?)+)", text, re.IGNORECASE)
        if not m:
            return []
        return [l.lstrip("-•* \t").strip() for l in m.group(1).split("\n") if l.strip() and l.strip()[0] in "-•*"]

    result["adoption_score"] = extract_num("ADOPTION_SCORE")
    result["sentiment"]      = extract_line("SENTIMENT_BREAKDOWN")

    vm = re.search(r"OVERALL_VERDICT\s*:\s*(.+?)(?=\n[A-Z_]{3,}\s*:|\Z)", text, re.IGNORECASE | re.DOTALL)
    result["verdict"] = vm.group(1).strip() if vm else "N/A"

    result["pain_points"] = extract_bullets("TOP_PAIN_POINTS", "TOP_STRENGTHS")
    result["strengths"]   = extract_bullets("TOP_STRENGTHS", "KEY_OBJECTIONS")
    result["objections"]  = extract_bullets("KEY_OBJECTIONS", "PRIORITY_SCORES")

    # Priority scores — explicitly bounded between PRIORITY_SCORES and PERSONA_HONESTY
    result["priority_scores"] = {}
    pm = re.search(r"PRIORITY_SCORES\s*:.*?\n(.*?)(?=PERSONA_HONESTY\s*:|$)", text, re.IGNORECASE | re.DOTALL)
    if pm:
        for line in pm.group(1).split("\n"):
            mm = re.match(r"\s*[-•*]\s*(.+?):\s*\[?(\d+)\]?", line)
            if mm:
                result["priority_scores"][mm.group(1).strip()] = int(mm.group(2))

    # Persona honesty — explicitly bounded between PERSONA_HONESTY and OVERALL_VERDICT
    result["honesty_scores"] = {}
    hm = re.search(r"PERSONA_HONESTY\s*:.*?\n(.*?)(?=OVERALL_VERDICT\s*:|$)", text, re.IGNORECASE | re.DOTALL)
    if hm:
        for line in hm.group(1).split("\n"):
            mm = re.match(r"\s*[-•*]\s*(.+?):\s*\[?(\d+)\]?", line)
            if mm:
                result["honesty_scores"][mm.group(1).strip()] = int(mm.group(2))

    return result


def parse_authenticity(text):
    result = {}

    def extract_num(key):
        m = re.search(rf"{key}[^\d]*(\d+)", text, re.IGNORECASE)
        return m.group(1).strip() if m else "N/A"

    def extract_line(key):
        m = re.search(rf"{key}\s*:\s*(.+?)$", text, re.IGNORECASE | re.MULTILINE)
        return m.group(1).strip() if m else "N/A"

    result["score"]            = extract_num("AUTHENTICITY_SCORE")
    result["groupthink"]       = extract_line("GROUPTHINK_DETECTED")
    result["groupthink_notes"] = extract_line("GROUPTHINK_NOTES")
    result["sycophancy_risk"]  = extract_line("CORPORATE_SYCOPHANCY_RISK")

    vm = re.search(r"FILTER_VERDICT\s*:\s*(.+?)(?=\n[A-Z_]{3,}\s*:|\Z)", text, re.IGNORECASE | re.DOTALL)
    result["verdict"] = vm.group(1).strip() if vm else "N/A"

    om = re.search(r"OUTLIER_VOICES\s*:.*?\n((?:\s*[-•*].+\n?)+)", text, re.IGNORECASE)
    result["outliers"] = []
    if om:
        result["outliers"] = [l.lstrip("-•* \t").strip() for l in om.group(1).split("\n") if l.strip() and l.strip()[0] in "-•*"]

    return result


def parse_next_steps(text):
    steps = []
    for i in range(1, 4):
        tm = re.search(rf"PANEL_{i}_TITLE:\s*(.+?)(?:\n|$)", text, re.IGNORECASE)
        rm = re.search(rf"PANEL_{i}_RATIONALE:\s*(.+?)(?:\n|$)", text, re.IGNORECASE)
        pm = re.search(rf"PANEL_{i}_PROFILE:\s*(.+?)(?:\n|$)", text, re.IGNORECASE)
        if tm:
            steps.append({
                "title": tm.group(1).strip(),
                "rationale": rm.group(1).strip() if rm else "",
                "profile": pm.group(1).strip() if pm else "",
            })
    return steps


# ─────────────────────────────────────────────
# RENDER HELPERS
# ─────────────────────────────────────────────
def bar_html(value, color_class="auto"):
    try:
        pct = min(max(int(value), 0), 100)
    except Exception:
        pct = 0
    if color_class == "auto":
        if pct >= 60:   color_class = "ga-bar-fill-green"
        elif pct >= 40: color_class = "ga-bar-fill"
        elif pct >= 25: color_class = "ga-bar-fill-warn"
        else:           color_class = "ga-bar-fill-danger"
    return f'<div class="ga-bar-track"><div class="{color_class}" style="width:{pct}%"></div></div>'


def honesty_chip(score):
    if score >= 75:   return f'<span class="honesty-chip hc-high">{score}</span>'
    elif score >= 50: return f'<span class="honesty-chip hc-mid">{score}</span>'
    else:             return f'<span class="honesty-chip hc-low">{score}</span>'


def verdict_badge(adoption):
    try: a = int(adoption)
    except Exception: return '<span class="verdict-badge verdict-yellow">—</span>'
    if a >= 75:   return f'<span class="verdict-badge verdict-green">🟢 Strong Buy Signal</span>'
    elif a >= 55: return f'<span class="verdict-badge verdict-yellow">🟡 Conditional Interest</span>'
    elif a >= 35: return f'<span class="verdict-badge verdict-orange">🟠 Skeptical — Needs Work</span>'
    else:         return f'<span class="verdict-badge verdict-red">🔴 Significant Resistance</span>'


def scorecard_delta(value):
    try: v = int(value)
    except Exception: return ""
    if v >= 70:   return f'<div class="ga-scorecard-delta delta-green">▲ Above threshold</div>'
    elif v >= 45: return f'<div class="ga-scorecard-delta delta-yellow">◆ Borderline</div>'
    else:         return f'<div class="ga-scorecard-delta delta-red">▼ Below threshold</div>'


def tip(text):
    """Returns an inline CSS tooltip '?' icon with hover text."""
    return f'<span class="tip-wrap"><span class="tip-icon">?</span><span class="tip-text">{text}</span></span>'


# ─────────────────────────────────────────────
# DOCX EXPORT
# ─────────────────────────────────────────────
def build_export_docx(pitch, personas, priorities, transcript_lines, insights_text,
                       auth, next_steps, focus_guide, phase2_messages):
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
        return p

    def body(text):
        doc.add_paragraph(text)

    def rule():
        doc.add_paragraph("─" * 60)

    # Title
    title = doc.add_heading("SimGroupAI — Simulation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    doc.add_paragraph(f"Panel: {', '.join([p['name'] for p in personas])}")
    doc.add_paragraph(f"Priorities: {', '.join(priorities) if priorities else 'None specified'}")
    doc.add_paragraph("")

    # Pitch
    heading("Pitch", 1)
    body(pitch)

    # Key Scores
    heading("Key Scores", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Score"
    scores = [
        ("Adoption Score", insights_text),
        ("Authenticity Score", auth.get("score", "N/A")),
        ("Groupthink Detected", auth.get("groupthink", "N/A")),
        ("Sycophancy Risk", auth.get("sycophancy_risk", "N/A")),
    ]
    # parse adoption from insights_text
    am = re.search(r"ADOPTION_SCORE[^\d]*(\d+)", insights_text, re.IGNORECASE)
    adoption_val = am.group(1) if am else "N/A"
    sm = re.search(r"SENTIMENT_BREAKDOWN\s*:\s*(.+?)$", insights_text, re.IGNORECASE | re.MULTILINE)
    sentiment_val = sm.group(1).strip() if sm else "N/A"
    for label, val in [
        ("Adoption Score", adoption_val),
        ("Sentiment", sentiment_val),
        ("Authenticity Score", auth.get("score", "N/A")),
        ("Groupthink Detected", auth.get("groupthink", "N/A")),
        ("Sycophancy Risk", auth.get("sycophancy_risk", "N/A")),
    ]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(val)

    doc.add_paragraph("")

    # Insights from text
    heading("Key Findings", 1)
    for section in ["TOP_PAIN_POINTS", "TOP_STRENGTHS", "KEY_OBJECTIONS"]:
        label = section.replace("_", " ").title()
        heading(label, 2)
        m = re.search(rf"{section}\s*:.*?\n((?:\s*[-•*].+\n?)+)", insights_text, re.IGNORECASE)
        if m:
            for line in m.group(1).split("\n"):
                l = line.lstrip("-•* \t").strip()
                if l:
                    doc.add_paragraph(l, style="List Bullet")

    # Priority scores
    pm = re.search(r"PRIORITY_SCORES\s*:.*?\n(.*?)(?=PERSONA_HONESTY\s*:|$)", insights_text, re.IGNORECASE | re.DOTALL)
    if pm:
        heading("Priority Dimension Scores", 2)
        for line in pm.group(1).split("\n"):
            mm = re.match(r"\s*[-•*]\s*(.+?):\s*\[?(\d+)\]?", line)
            if mm:
                doc.add_paragraph(f"{mm.group(1).strip()}: {mm.group(2)}/100", style="List Bullet")

    # Honesty scores
    hm = re.search(r"PERSONA_HONESTY\s*:.*?\n(.*?)(?=OVERALL_VERDICT\s*:|$)", insights_text, re.IGNORECASE | re.DOTALL)
    if hm:
        heading("Persona Honesty Scores", 2)
        for line in hm.group(1).split("\n"):
            mm = re.match(r"\s*[-•*]\s*(.+?):\s*\[?(\d+)\]?", line)
            if mm:
                doc.add_paragraph(f"{mm.group(1).strip()}: {mm.group(2)}/100", style="List Bullet")

    # Verdict
    vm = re.search(r"OVERALL_VERDICT\s*:\s*(.+?)(?=\n[A-Z_]{3,}\s*:|\Z)", insights_text, re.IGNORECASE | re.DOTALL)
    if vm:
        heading("Overall Verdict", 2)
        body(vm.group(1).strip())

    # Next Steps
    if next_steps:
        heading("Recommended Next Research Steps", 1)
        for i, s in enumerate(next_steps, 1):
            heading(f"{i}. {s['title']}", 2)
            body(s["rationale"])
            body(f"Target profile: {s['profile']}")

    # Focus Guide
    if focus_guide:
        heading("Human Focus Group Guide", 1)
        body(focus_guide)

    # Transcript
    heading("Focus Group Transcript", 1)
    for item in transcript_lines:
        p = doc.add_paragraph()
        run = p.add_run(f"{item['speaker']}: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x8a, 0xb4, 0xf8)
        p.add_run(item["message"])

    # Phase 2
    if phase2_messages:
        heading("Phase 2 — Follow-Up Conversation", 1)
        for msg in phase2_messages:
            if msg["role"] == "user":
                p = doc.add_paragraph()
                r = p.add_run("Moderator: ")
                r.bold = True
                r.font.color.rgb = RGBColor(0xfb, 0xbc, 0x04)
                p.add_run(msg["content"])
            else:
                for line in msg["content"].strip().split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    mm = re.match(r"^\[?([^\]:\n]+)\]?:\s*(.+)$", line)
                    if mm:
                        p = doc.add_paragraph()
                        r = p.add_run(f"{mm.group(1).strip()}: ")
                        r.bold = True
                        p.add_run(mm.group(2).strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────
def render_persona_card_sidebar(p):
    """Renders expanded persona profile content inside a sidebar expander."""
    vals = ", ".join(p.get("core_values", []))
    st.markdown(
        f"""<div style="font-size:0.78rem; line-height:1.75; color:#bdc1c6; padding: 4px 0;">
<b style="color:#e8eaed;">📍 Location:</b> {p.get('location_type', '—')}<br>
<b style="color:#e8eaed;">💰 Income:</b> {p.get('income_bracket', '—')}<br>
<b style="color:#e8eaed;">⭐ Values:</b> {vals}<br>
<b style="color:#e8eaed;">🎭 Rating Style:</b> {p.get('rating_style', 'Balanced')}<br>
<b style="color:#e8eaed;">🗂 Background:</b> {p.get('background', '—')}<br>
<b style="color:#e8eaed;">💬 Communication:</b> {p.get('communication_style', '—')}
</div>""",
        unsafe_allow_html=True,
    )


def render_sidebar(all_personas):
    with st.sidebar:
        st.markdown("### 🧠 SimGroupAI")
        st.markdown("---")
        st.markdown("**API Key**")
        key_input = st.text_input("OpenAI API Key", type="password",
                                   value=st.session_state.api_key,
                                   placeholder="sk-...", label_visibility="collapsed")
        if key_input:
            st.session_state.api_key = key_input

        st.markdown("---")
        st.markdown("**Panel**")
        panel_mode = st.radio("Panel type", ["Default Library", "Custom AI Panel"],
                               index=0 if st.session_state.panel_mode == "library" else 1,
                               label_visibility="collapsed")
        st.session_state.panel_mode = "library" if panel_mode == "Default Library" else "custom"

        selected_personas = []

        if st.session_state.panel_mode == "library":
            for p in all_personas:
                checked = st.checkbox(f"{p['icon']} {p['name']} · {p['age_group']}", value=True, key=f"persona_{p['name']}")
                if checked:
                    selected_personas.append(p)
                with st.expander(f"  ↳ {p['name']} profile", expanded=False):
                    render_persona_card_sidebar(p)

            if len(selected_personas) >= 2:
                with st.expander("💾 Save this panel"):
                    pname = st.text_input("Name", placeholder="e.g. Skeptics", key="save_pname")
                    if st.button("Save") and pname:
                        st.session_state.saved_panels[pname] = list(selected_personas)
                        st.success(f"Saved: {pname}")

            if st.session_state.saved_panels:
                with st.expander("📂 Load saved panel"):
                    sname = st.selectbox("Panel", list(st.session_state.saved_panels.keys()))
                    if st.button("Load"):
                        st.session_state.custom_panel = st.session_state.saved_panels[sname]
                        st.session_state.panel_mode = "custom"
                        st.rerun()
        else:
            if st.session_state.custom_panel:
                for p in st.session_state.custom_panel:
                    st.markdown(f"{p['icon']} **{p['name']}** · {p['age_group']}")
                    with st.expander(f"  ↳ {p['name']} profile", expanded=False):
                        render_persona_card_sidebar(p)
                selected_personas = st.session_state.custom_panel
                if st.button("✕ Clear panel"):
                    st.session_state.custom_panel = None
                    st.session_state.panel_mode = "library"
                    st.rerun()
            else:
                audience_desc = st.text_area("Describe target audience",
                    placeholder="e.g. Health-conscious millennials in urban areas...",
                    height=90, label_visibility="collapsed", key="audience_desc")
                if st.button("⚡ Generate Panel"):
                    client = get_client()
                    if not client:
                        st.error("Enter API key first.")
                    elif not audience_desc.strip():
                        st.warning("Describe your audience.")
                    else:
                        with st.spinner("Building panel..."):
                            try:
                                ps = generate_custom_panel(client, audience_desc)
                                st.session_state.custom_panel = ps
                                st.rerun()
                            except Exception as e:
                                st.error(f"Error: {e}")

        st.markdown("---")
        st.markdown("**Priorities**")
        st.caption("Select what matters most for this product:")
        selected_priorities = [p for p in PRIORITIES if st.checkbox(p, value=False, key=f"prio_{p}")]
        custom_p = st.text_input("Custom priority", placeholder="e.g. Gamification", key="custom_prio")
        if custom_p.strip():
            selected_priorities.append(custom_p.strip())

        return selected_personas, selected_priorities


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def main():
    init_session_state()
    inject_css()

    st.markdown("""
    <div class="ga-topbar">
        <div class="ga-logo"><span class="ga-logo-dot"></span>SimGroupAI</div>
        <div class="ga-property">Consumer Intelligence Platform</div>
    </div>
    """, unsafe_allow_html=True)

    all_personas = load_personas()
    selected_personas, selected_priorities = render_sidebar(all_personas)

    pitch = st.text_area("Product or Service Pitch",
        placeholder='Describe your product or service. e.g. "WealthMind is a personal finance AI that builds custom savings plans — no spreadsheets required."',
        height=100)

    col_btn, col_info = st.columns([1, 5])
    with col_btn:
        run_btn = st.button("▶  Run Simulation", use_container_width=True)
    with col_info:
        if selected_personas:
            st.caption(f"{len(selected_personas)} panelists · {len(selected_priorities)} priorities selected")
        else:
            st.caption("Select at least 2 personas in the sidebar.")

    # ── Run ──
    if run_btn:
        client = get_client()
        if not client:
            st.error("Enter your OpenAI API key in the sidebar.")
        elif not pitch.strip():
            st.warning("Enter a pitch first.")
        elif len(selected_personas) < 2:
            st.warning("Select at least 2 personas.")
        else:
            st.session_state.simulation_done = False
            st.session_state.phase2_messages = []
            st.session_state.active_view = "phase1"
            prog = st.progress(0, text="Initializing...")
            try:
                prog.progress(10, text="🗣  Generating transcript...")
                raw = generate_simulation(client, pitch, selected_personas, selected_priorities)
                st.session_state.transcript_lines = parse_transcript(raw, selected_personas)
                st.session_state.raw_transcript = raw

                prog.progress(35, text="📊  Extracting insights...")
                it = generate_insights(client, pitch, raw, selected_personas, selected_priorities)
                st.session_state.insights = parse_insights(it)
                st.session_state.insights_text = it

                prog.progress(60, text="🔍  Authenticity filter...")
                at = check_authenticity(client, pitch, raw)
                st.session_state.auth = parse_authenticity(at)

                prog.progress(75, text="🗺  Building focus guide...")
                st.session_state.focus_guide = generate_focus_guide(client, pitch, raw, it)

                prog.progress(88, text="🔭  Next-step recommendations...")
                nst = generate_next_steps(client, pitch, it, selected_personas)
                st.session_state.next_steps = parse_next_steps(nst)

                prog.progress(100, text="✅  Done.")
                st.session_state.simulation_done = True
                st.session_state.selected_personas = selected_personas
                st.session_state.pitch = pitch
                st.session_state.selected_priorities = selected_priorities

            except Exception as e:
                prog.empty()
                st.error(f"Error: {e}")

    # ── Results ──
    if st.session_state.simulation_done:
        insights   = st.session_state.insights
        auth       = st.session_state.auth
        personas   = st.session_state.selected_personas
        priorities = st.session_state.selected_priorities
        honesty    = insights.get("honesty_scores", {})
        adoption   = insights.get("adoption_score", "—")
        auth_score = auth.get("score", "—")

        # Scorecard row
        st.markdown(f"""
        <div class="ga-scorecard-row">
            <div class="ga-scorecard">
                <div class="ga-scorecard-label">Adoption Score {tip("The likelihood this panel would adopt or purchase the product, scored 0–100. Above 70 = strong signal.")}</div>
                <div class="ga-scorecard-value">{adoption}<span style="font-size:1rem;color:#5f6368">/100</span></div>
                {scorecard_delta(adoption)}
            </div>
            <div class="ga-scorecard">
                <div class="ga-scorecard-label">Authenticity Score {tip("How candid and unfiltered the discussion was. Low scores indicate groupthink or sycophantic responses, scored 0–100.")}</div>
                <div class="ga-scorecard-value">{auth_score}<span style="font-size:1rem;color:#5f6368">/100</span></div>
                {scorecard_delta(auth_score)}
            </div>
            <div class="ga-scorecard">
                <div class="ga-scorecard-label">Sentiment</div>
                <div style="font-size:0.88rem;color:#bdc1c6;margin-top:6px;line-height:1.4;">{insights.get("sentiment","—")}</div>
            </div>
            <div class="ga-scorecard" style="flex:2">
                <div class="ga-scorecard-label">Verdict</div>
                <div style="margin-top:6px;">{verdict_badge(adoption)}</div>
                <div style="font-size:0.78rem;color:#9aa0a6;margin-top:6px;line-height:1.4;">{insights.get("verdict","")}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # Phase navigation — session-state buttons (persists across reruns, unlike st.tabs)
        col_p1, col_p2, col_spacer = st.columns([1, 1, 4])
        with col_p1:
            if st.button("📊  Phase 1 — Results",
                          type="primary" if st.session_state.active_view == "phase1" else "secondary",
                          use_container_width=True, key="nav_p1"):
                st.session_state.active_view = "phase1"
                st.rerun()
        with col_p2:
            if st.button("💬  Phase 2 — Follow-Up",
                          type="primary" if st.session_state.active_view == "phase2" else "secondary",
                          use_container_width=True, key="nav_p2"):
                st.session_state.active_view = "phase2"
                st.rerun()

        st.markdown("---")

        # ══════════════════════════════════════════
        # PHASE 1
        # ══════════════════════════════════════════
        if st.session_state.active_view == "phase1":

            # Exports at top
            st.markdown('<div class="ga-section-title">Export</div>', unsafe_allow_html=True)
            ex1, ex2, _ = st.columns([1, 1, 4])
            with ex1:
                try:
                    docx_bytes = build_export_docx(
                        st.session_state.pitch, personas, priorities,
                        st.session_state.transcript_lines, st.session_state.insights_text,
                        auth, st.session_state.next_steps,
                        st.session_state.focus_guide, st.session_state.phase2_messages
                    )
                    st.download_button("⬇  Full Report (.docx)", data=docx_bytes,
                                        file_name="simgroup_report.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="dl_full")
                except Exception as e:
                    st.error(f"Export error: {e}")
            with ex2:
                st.download_button("⬇  Focus Guide (.docx)",
                                    data=build_focus_guide_docx(st.session_state.focus_guide, st.session_state.pitch),
                                    file_name="simgroup_focus_guide.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="dl_guide")

            # Meet the Panel
            st.markdown('<div class="ga-section-title">Meet the Panel</div>', unsafe_allow_html=True)
            panel_cols = st.columns(len(personas))
            for idx, p in enumerate(personas):
                with panel_cols[idx]:
                    vals_short = " · ".join(p.get("core_values", [])[:2])
                    st.markdown(
                        f"""<div style="background:#35363a; border:1px solid #3c3c3f; border-radius:8px;
                        padding:10px 12px; text-align:center; margin-bottom:4px;">
                        <div style="font-size:1.6rem;">{p.get('icon','👤')}</div>
                        <div style="font-size:0.8rem; font-weight:600; color:#e8eaed; margin:3px 0;">{p['name']}</div>
                        <div style="font-size:0.68rem; color:#9aa0a6;">{p.get('age_group','')}</div>
                        <div style="font-size:0.68rem; color:#9aa0a6; margin-top:2px;">{p.get('income_bracket','')}</div>
                        </div>""",
                        unsafe_allow_html=True,
                    )
                    with st.expander("Profile", expanded=False):
                        render_persona_card_sidebar(p)

            # Dashboard
            col_left, col_right = st.columns(2)
            with col_left:
                st.markdown(f'<div class="ga-section-title">Scores</div>', unsafe_allow_html=True)
                for label, val, tooltip_text in [
                    ("Adoption Score", adoption, "Likelihood of adoption, 0–100."),
                    ("Authenticity Score", auth_score, "How candid the discussion was, 0–100."),
                ]:
                    try: v = int(val)
                    except Exception: v = 0
                    st.markdown(f"""
                    <div class="ga-bar-row">
                        <div class="ga-bar-label-row"><span>{label}</span><span style="color:#8ab4f8;font-weight:500">{v}</span></div>
                        {bar_html(v)}
                    </div>""", unsafe_allow_html=True)

                priority_scores = insights.get("priority_scores", {})
                if priority_scores:
                    st.markdown(
                        f'<div class="ga-section-title">Priority Dimension Scores {tip("How strongly the panel engaged with each priority you selected, scored 0–100. Higher = more discussed, validated, or contested.")}</div>',
                        unsafe_allow_html=True)
                    for dim, score in priority_scores.items():
                        st.markdown(f"""
                        <div class="ga-bar-row">
                            <div class="ga-bar-label-row"><span>{dim}</span><span style="color:#8ab4f8;font-weight:500">{score}</span></div>
                            {bar_html(score)}
                        </div>""", unsafe_allow_html=True)

                if honesty:
                    st.markdown(
                        f'<div class="ga-section-title">Persona Honesty Scores {tip("How candid and unfiltered each individual persona was during the discussion, scored 0–100. Higher = more direct and critical.")}</div>',
                        unsafe_allow_html=True)
                    for name, score in honesty.items():
                        p_obj = next((p for p in personas if p["name"] == name), None)
                        icon = p_obj["icon"] if p_obj else "👤"
                        st.markdown(f"""
                        <div class="ga-bar-row">
                            <div class="ga-bar-label-row"><span>{icon} {name}</span>{honesty_chip(score)}</div>
                            {bar_html(score)}
                        </div>""", unsafe_allow_html=True)

            with col_right:
                st.markdown('<div class="ga-section-title">Key Findings</div>', unsafe_allow_html=True)
                for items, dot, label in [
                    (insights.get("pain_points",[]), "dot-red",    "Pain Points"),
                    (insights.get("strengths",[]),   "dot-green",  "Strengths"),
                    (insights.get("objections",[]),  "dot-yellow", "Key Objections"),
                ]:
                    if items:
                        st.markdown(f'<div style="font-size:0.72rem;font-weight:500;color:#9aa0a6;text-transform:uppercase;letter-spacing:0.7px;margin-bottom:6px;margin-top:12px;">{label}</div>', unsafe_allow_html=True)
                        for item in items:
                            st.markdown(f'<div class="ga-finding-row"><div class="ga-finding-dot {dot}"></div><span>{item}</span></div>', unsafe_allow_html=True)

                gnotes = auth.get("groupthink_notes", "")
                if gnotes and gnotes.lower() not in ("none", "n/a", ""):
                    st.markdown(f"""
                    <div style="margin-top:16px;padding:10px 14px;background:#3a2810;border:1px solid #f57c00;border-radius:6px;font-size:0.8rem;color:#ffb74d;line-height:1.45;">
                    ⚠ <strong>Groupthink Note:</strong> {gnotes}</div>""", unsafe_allow_html=True)

            # Next Steps
            if st.session_state.next_steps:
                st.markdown('<div class="ga-section-title">Recommended Next Research Steps</div>', unsafe_allow_html=True)
                ns_cols = st.columns(len(st.session_state.next_steps))
                for col, step in zip(ns_cols, st.session_state.next_steps):
                    with col:
                        st.markdown(f"""
                        <div class="ns-card">
                            <div class="ns-title">→ {step["title"]}</div>
                            <div class="ns-body">{step["rationale"]}</div>
                            <div class="ns-profile">Target: {step["profile"]}</div>
                        </div>""", unsafe_allow_html=True)

            # Transcript
            st.markdown('<div class="ga-section-title">Focus Group Transcript</div>', unsafe_allow_html=True)
            for item in st.session_state.transcript_lines:
                speaker = item["speaker"]
                persona = item.get("persona")
                icon = persona["icon"] if persona else "👤"
                hs = honesty.get(speaker, None)
                chip = honesty_chip(hs) if hs is not None else ""
                st.markdown(f"""
                <div class="tx-bubble">
                    <div class="tx-name">{icon} {speaker}{chip}</div>
                    {item["message"]}
                </div>""", unsafe_allow_html=True)

        # ══════════════════════════════════════════
        # PHASE 2
        # ══════════════════════════════════════════
        else:
            # Export at top of Phase 2
            st.markdown('<div class="ga-section-title">Export</div>', unsafe_allow_html=True)
            ex_col, _ = st.columns([1, 5])
            with ex_col:
                try:
                    docx_bytes = build_export_docx(
                        st.session_state.pitch, personas, priorities,
                        st.session_state.transcript_lines, st.session_state.insights_text,
                        auth, st.session_state.next_steps,
                        st.session_state.focus_guide, st.session_state.phase2_messages
                    )
                    st.download_button("⬇  Full Report (.docx)", data=docx_bytes,
                                        file_name="simgroup_report_with_followup.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="dl_full_p2")
                except Exception as e:
                    st.error(f"Export error: {e}")

            st.markdown("""
            <div class="p2-info">
                💬 <strong>Phase 2 — Follow-Up Q&A</strong><br>
                Ask follow-up questions directly to the panelists. Address the group or call out a specific participant by name.
                Press <strong>Enter</strong> or click <strong>Send</strong> to submit.
            </div>""", unsafe_allow_html=True)

            # Conversation history
            for msg in st.session_state.phase2_messages:
                if msg["role"] == "user":
                    st.markdown(f"""
                    <div class="tx-bubble tx-bubble-user">
                        <div class="tx-name tx-name-user">🎤 You (Moderator)</div>
                        {msg["content"]}
                    </div>""", unsafe_allow_html=True)
                else:
                    for line in msg["content"].strip().split("\n"):
                        line = line.strip()
                        if not line: continue
                        mm = re.match(r"^\[?([^\]:\n]+)\]?:\s*(.+)$", line)
                        if mm:
                            speaker = mm.group(1).strip()
                            text    = mm.group(2).strip()
                            p_obj   = next((p for p in personas if p["name"].lower() in speaker.lower()), None)
                            icon    = p_obj["icon"] if p_obj else "👤"
                            hs      = honesty.get(speaker, None)
                            chip    = honesty_chip(hs) if hs is not None else ""
                            st.markdown(f"""
                            <div class="tx-bubble">
                                <div class="tx-name">{icon} {speaker}{chip}</div>
                                {text}
                            </div>""", unsafe_allow_html=True)

            # ── Phase 2 input — st.form enables Enter key and clears on submit ──
            with st.form(key="phase2_form", clear_on_submit=True):
                user_q = st.text_input(
                    "Your question",
                    placeholder='e.g. "What would make you actually buy this?" or "Robert, what\'s your biggest concern?"',
                    label_visibility="collapsed",
                )
                submitted = st.form_submit_button("Send →", use_container_width=False)

            if submitted and user_q.strip():
                client = get_client()
                if not client:
                    st.error("API key required.")
                else:
                    st.session_state.phase2_messages.append({"role": "user", "content": user_q})
                    st.session_state.active_view = "phase2"  # stay on phase 2
                    with st.spinner("Panelists responding..."):
                        try:
                            response = generate_followup(
                                client, st.session_state.pitch,
                                st.session_state.raw_transcript, personas,
                                st.session_state.phase2_messages, user_q,
                            )
                            st.session_state.phase2_messages.append({
                                "role": "assistant", "speaker": "panel", "content": response,
                            })
                        except Exception as e:
                            st.error(f"Error: {e}")
                    st.rerun()

    else:
        st.markdown("""
        <div style="text-align:center;padding:70px 20px;">
            <div style="font-size:2.5rem;margin-bottom:14px;">🧬</div>
            <div style="font-size:1rem;font-weight:500;color:#5f6368;margin-bottom:8px;">No simulation running</div>
            <div style="font-size:0.84rem;color:#3c3c3f;">Enter your pitch above and click <strong style="color:#8ab4f8">Run Simulation</strong></div>
        </div>""", unsafe_allow_html=True)


def build_focus_guide_docx(focus_guide, pitch):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    title = doc.add_heading("SimGroupAI — Human Focus Group Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Pitch: {pitch}")
    doc.add_paragraph("")
    doc.add_paragraph(focus_guide)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


if __name__ == "__main__":
    main()
